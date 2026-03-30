import json
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from google import genai
from dotenv import load_dotenv
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import cm
from google.genai import types

from xml.sax.saxutils import escape as xml_escape
import re
import os
import time
import traceback

_log_lock = threading.Lock()


def _write_log(log_path, msg):
    """スレッドセーフなログ書き込み"""
    with _log_lock:
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(msg)

# ============================================
# 設定
# ============================================
load_dotenv()

WARS_ID = os.getenv("WARS_ID", "")
#FONT_PATH = "C:/Windows/Fonts/meiryo.ttc"
FONT_NAME = "meiryo"
THINKING_MODELS = {"gemini-3.1-pro-preview", "gemini-3.1-flash-preview"}

# 次回のGemini呼び出しで使い始めるAPIキーのインデックス（0〜4、周回管理）
_key_index = 0

# プロンプト
PROMPT_PREFIX = f"""
あなたは将棋の対局解説を行うプロ棋士AIです。
これからあなたに、私(={WARS_ID})が負けた対局の棋譜データをこのプロンプトの末尾に提示します。
この対局における
1. 対局の流れ・総評
2. 私の敗因
  - この項目はどの局面について解説しているかを分かるように、必ず各段落の前に※※〇〇手目の局面※※という1行を書いてください
について、それぞれ教えてください
"""

def delete_png_and_docx_files(root_dir):
    """
    指定されたディレクトリ配下のすべての .png および .docx ファイルを削除する。
    ただし、banmen_all_full.png は削除対象から除外する。
    """
    for dirpath, _, filenames in os.walk(root_dir):
        for filename in filenames:
            ext = filename.lower()
            if ext.endswith('.png') or ext.endswith('.docx'):
                # banmen_all_full.png は残す
                if filename == "banmen_all_full.png":
                    continue
                
                file_path = os.path.join(dirpath, filename)
                try:
                    os.remove(file_path)
                    print(f"削除しました: {file_path}")
                except Exception as e:
                    print(f"削除できませんでした: {file_path} 理由: {e}")

def cut_unnecessary_kif(ROOT, log_path, kif_path):
    entity_kif = ""
    header_keyword = "手数----指手---------消費時間--"
    resign_keywords = ["投了", "詰み", "切れ負け", "千日手"] # 千日手も念のため追加

    header_line_num = None
    resign_line_num = None

    lines = []
    # エンコーディングエラー回避のため errors='ignore' や 'replace' を検討しても良いが、一旦utf-8
    with open(kif_path, encoding="utf-8") as f:
        for i, line in enumerate(f, start=1):
            stripped = line.rstrip("\n")
            lines.append(stripped)

            if header_keyword in stripped:
                header_line_num = i

            if any(keyword in stripped for keyword in resign_keywords):
                resign_line_num = i

    if header_line_num is None:
        # ヘッダーが見つからない場合は処理できないため空文字等を返すかエラーにする
        # ここではエラーを出さずにスキップする運用を想定し空文字を返す例
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"[Warn] ヘッダーが見つかりません: {kif_path}\n")
        print(f" [Warn] ヘッダーが見つかりません: {kif_path}")
        return ""
    
    if resign_line_num is None:
        # 投了等が見つからない場合はファイルの最後まで読む
        resign_line_num = len(lines) + 1

    kif_data = lines[header_line_num:resign_line_num]
    for i, d in enumerate(kif_data):
        # 最終行以外は改行を入れる
        if i < len(kif_data) - 1:
            entity_kif += d[:13].strip() + "\n"
        else:
            entity_kif += d[:13].strip()

    return entity_kif


def call_gemini_analysis(raw_kif, log_path):
    """Gemini APIを呼び出して解析結果を返す（GEMINI_MODELSを順にトライ、API_KEY01〜05をローテーション）"""
    global _key_index

    prompt = f"{PROMPT_PREFIX}\n{raw_kif}"

    gemini_models_str = os.getenv("GEMINI_MODELS", "gemini-3.1-pro-preview,gemini-3.1-flash-preview,gemini-2.5-pro,gemini-2.5-flash")
    model_list = [m.strip() for m in gemini_models_str.split(",") if m.strip()]

    api_keys = [
        os.getenv("API_KEY01"),
        os.getenv("API_KEY02"),
        os.getenv("API_KEY03"),
        os.getenv("API_KEY04"),
        os.getenv("API_KEY05"),
    ]
    num_keys = len(api_keys)

    last_error = None
    for model_idx, model in enumerate(model_list):
        use_thinking = model in THINKING_MODELS
        for attempt in range(num_keys):
            idx = (_key_index + attempt) % num_keys
            key_num = f"{idx + 1:02d}"
            api_key = api_keys[idx]

            if not api_key:
                msg = f"[WARN] API_KEY{key_num} が未設定のためスキップ\n"
                print(msg.strip())
                _write_log(log_path, msg)
                continue

            try:
                msg = f"[*] Gemini呼び出し (KEY{key_num}): {model}\n"
                print(msg.strip())
                _write_log(log_path, msg)
                client = genai.Client(api_key=api_key)
                if use_thinking:
                    config = types.GenerateContentConfig(
                        thinking_config=types.ThinkingConfig(
                            include_thoughts=True,
                            thinking_level="high"
                        )
                    )
                else:
                    config = types.GenerateContentConfig()
                response = client.models.generate_content(
                    model=model,
                    contents=prompt,
                    config=config
                )
                if response and response.text:
                    _key_index = (idx + 1) % num_keys
                    return response.text.strip()
            except Exception as e:
                last_error = e
                msg = f"[WARN] KEY{key_num}/{model} 失敗: {e}\n"
                print(msg.strip())
                _write_log(log_path, msg)
                if attempt < num_keys - 1:
                    time.sleep(10)

        # 全キー失敗 → 次のモデルへ
        if model_idx < len(model_list) - 1:
            msg = f"[WARN] モデル {model} 全キー失敗。次のモデルへ移行します\n"
            print(msg.strip())
            _write_log(log_path, msg)
            time.sleep(10)

    raise RuntimeError(f"Gemini API 呼び出し失敗（全モデル・全キー）: {last_error}")


def txt_to_docx_with_images(kif_text, response):
    """
    Geminiのレスポンスから「※※〇〇手目の局面※※」を抽出し、
    貼り付ける手数リストと、各手数までの棋譜データリストを返す。

    Returns:
        haru_tesuu: ["all", 手数1, 手数2, ...]  ("all"はfull kif)
        kif_list:   [full_kif, kif_up_to_手数1, kif_up_to_手数2, ...]
    """
    lines = kif_text.splitlines()

    # WARS_ID の先手/後手を判定
    teban = None
    for line in lines:
        if WARS_ID in line:
            if "先手" in line:
                teban = "先手"
            elif "後手" in line:
                teban = "後手"

    # レスポンスから「※※〇〇手目の局面※※」を抽出
    pattern = r"※※\s*\d+手目の局面\s*※※"
    matches = re.findall(pattern, response)

    # 自分の手番の1手前の局面図を貼るために手数を調整
    haru_tesuu = ["all"]
    for kyokumen in matches:
        extracted = re.search(r"※※(\d+)手目の局面※※", kyokumen)
        if extracted:
            number = int(extracted.group(1))
            if number <= 1:
                continue
            if number % 2 == 1 and teban == "先手":
                number -= 1
            if number % 2 == 0 and teban == "後手":
                number -= 1
            haru_tesuu.append(number)
            print(f"取り出した手数: {number}")
        else:
            print("手数が見つかりませんでした。")

    # haru_tesuu[0]="all" に対応する full kif を先頭に追加
    kif_list = [kif_text]

    # 各手数までの棋譜データを生成
    move_pattern = re.compile(r"^\s*(\d+)\s+([^\s]+)")
    for tesuu in haru_tesuu[1:]:  # "all" を除く
        kif_data = ""
        for line in lines:
            kif_data += line + "\n"
            m = move_pattern.match(line)
            if not m:
                continue
            move_number = int(m.group(1))
            if move_number == tesuu:
                kif_list.append(kif_data)
                break

    return haru_tesuu, kif_list, teban


def get_final_board_image(kif_text, save_prefix, teban):
    """
    UserLocal解析サイトに棋譜を送信し、盤面図を取得する。

    Args:
        kif_text:    送信する棋譜テキスト
        save_prefix: 保存先プレフィックス（例: /path/to/banmen_all）
                     → {save_prefix}_full.png, {save_prefix}_resize.png に保存

    Returns:
        analysis_url: 解析結果のURL
    """
    analysis_url = ""
    try:
        driver = webdriver.Chrome()
        driver.get("https://shogi-club-analysis.userlocal.jp/analysis")
        time.sleep(3)

        driver.find_element(By.CSS_SELECTOR, "button[data-target='#kifInputModal']").click()

        wait = WebDriverWait(driver, 10)
        textarea = wait.until(EC.visibility_of_element_located((By.NAME, "kif")))

        textarea.clear()
        textarea.send_keys(kif_text)

        driver.find_element(
            By.XPATH,
            "//button[@type='submit' and contains(text(), '更新')]"
        ).click()

        time.sleep(70)

        analysis_url = driver.current_url
        print(f"解析完了URL: {analysis_url}")

        full_path = f"{save_prefix}_full.png"
        resize_path = f"{save_prefix}_resize.png"

        driver.save_screenshot(full_path)

        image = Image.open(full_path)
        cropped = image.crop((0, 0, 350, 550))
        if teban == "後手":
            cropped = cropped.transpose(Image.FLIP_TOP_BOTTOM)
            cropped = cropped.transpose(Image.FLIP_LEFT_RIGHT)
        cropped.save(resize_path)

        print("=== 解析エンジン実行終了 ===")

    except Exception as e:
        print(f"盤面図取得エラー: {e}")
    finally:
        driver.quit()

    return analysis_url


def _register_font(ROOT):
    """日本語フォントを登録する（未登録の場合のみ）"""
    FONT_PATH = os.path.join(ROOT,"Fonts/meiryo.ttc")
    if FONT_NAME not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_PATH))


def _trim_response(response_txt):
    """「1. 対局の流れ・総評」より前の内容を削除する"""
    marker = "1. 対局の流れ・総評"
    idx = response_txt.find(marker)
    if idx != -1:
        return response_txt[idx:]
    return response_txt


def make_leport_docx(response_txt, battle_id, re_url, battle_dir, docx_path, all_kif_data, haru_tesuu):
    """レポートの.docxを作成する"""
    doc = Document()

    # 1ページ目: battle_id の振り返り + URL + 全体盤面図（full）
    header = f"{battle_id}の振り返り:\n{re_url}"
    doc.add_paragraph(header)

    fin_full_path = os.path.join(battle_dir, "banmen_all_full.png")
    if os.path.exists(fin_full_path):
        doc.add_picture(fin_full_path, width=Inches(5.0))

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # 「1. 対局の流れ・総評」より前の内容を削除
    trimmed = _trim_response(response_txt)

    # Geminiレスポンスを1行ずつ処理
    pattern_im = re.compile(r"※※\s*(\d+)手目の局面\s*※※")
    lines = trimmed.splitlines()
    # haru_tesuu[0]="all" に対応するカウンター（number<=1はスキップされているのでカウントしない）
    kyokumen_count = 0
    # 直前に改ページが入ったか（True=改ページ不要）
    # 初期値True: 1ページ目→PageBreak後の状態から開始
    last_was_page_break = True

    for line in lines:
        if line == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            last_was_page_break = True
            continue

        m = pattern_im.search(line)
        if m:
            number_in_response = int(m.group(1))
            if number_in_response <= 1:
                # txt_to_docx_with_images でスキップされた局面はカウントしない
                # last_was_page_break は変更しない（後続の最初の有効局面では改ページしない）
                doc.add_paragraph(line)
                continue
            kyokumen_count += 1
            # 直前に改ページがない場合のみ改ページ
            if not last_was_page_break:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            last_was_page_break = False
            doc.add_paragraph(line)
            # haru_tesuu[kyokumen_count] が調整済みの手数（haru_tesuu[0]="all"）
            if kyokumen_count < len(haru_tesuu):
                adjusted_tesuu = haru_tesuu[kyokumen_count]
            else:
                adjusted_tesuu = number_in_response
            img_path = os.path.join(battle_dir, f"banmen_{adjusted_tesuu}_resize.png")
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(2.5))
            continue

        doc.add_paragraph(line)

    # 最終ページ: 棋譜データ
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("棋譜データ")
    doc.add_paragraph(all_kif_data)

    doc.save(docx_path)
    print(f"docx保存: {docx_path}")


def make_report_pdf(ROOT,response_txt, battle_id, re_url, battle_dir, pdf_path, all_kif_data, haru_tesuu):
    """reportlab で PDF を直接生成する"""
    _register_font(ROOT)

    body_style = ParagraphStyle(
        "Body",
        fontName=FONT_NAME,
        fontSize=10,
        leading=14,
        wordWrap="CJK",
    )
    title_style = ParagraphStyle(
        "Title",
        fontName=FONT_NAME,
        fontSize=13,
        leading=20,
        wordWrap="CJK",
    )
    kif_style = ParagraphStyle(
        "Kif",
        fontName=FONT_NAME,
        fontSize=8,
        leading=12,
        wordWrap="CJK",
    )

    story = []

    # 1ページ目: タイトル + URL + 全体盤面図（full）
    story.append(Paragraph(xml_escape(f"{battle_id}の振り返り"), title_style))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(xml_escape(re_url), body_style))
    story.append(Spacer(1, 0.5 * cm))

    fin_full_path = os.path.join(battle_dir, "banmen_all_full.png")
    if os.path.exists(fin_full_path):
        story.append(RLImage(fin_full_path, width=17 * cm, height=11 * cm))
        #story.append(RLImage(fin_full_path, width=11 * cm, height=11 * cm))

    story.append(PageBreak())

    # 「1. 対局の流れ・総評」より前の内容を削除
    trimmed = _trim_response(response_txt)

    # Geminiレスポンスを1行ずつ処理
    pattern_im = re.compile(r"※※\s*(\d+)手目の局面\s*※※")
    lines = trimmed.splitlines()
    # haru_tesuu[0]="all" に対応するカウンター（number<=1はスキップされているのでカウントしない）
    kyokumen_count = 0
    # 直前に改ページが入ったか（True=改ページ不要）
    # 初期値True: 1ページ目→PageBreak後の状態から開始
    last_was_page_break = True

    for line in lines:
        if line == "---":
            story.append(PageBreak())
            last_was_page_break = True
            continue

        m = pattern_im.search(line)
        if m:
            number_in_response = int(m.group(1))
            if number_in_response <= 1:
                # txt_to_docx_with_images でスキップされた局面はカウントしない
                # last_was_page_break は変更しない（後続の最初の有効局面では改ページしない）
                story.append(Paragraph(xml_escape(line), body_style))
                continue
            kyokumen_count += 1
            # 直前に改ページがない場合のみ改ページ
            if not last_was_page_break:
                story.append(PageBreak())
            last_was_page_break = False
            story.append(Paragraph(xml_escape(line), body_style))
            story.append(Spacer(1, 0.3 * cm))
            # haru_tesuu[kyokumen_count] が調整済みの手数（haru_tesuu[0]="all"）
            if kyokumen_count < len(haru_tesuu):
                adjusted_tesuu = haru_tesuu[kyokumen_count]
            else:
                adjusted_tesuu = number_in_response
            img_path = os.path.join(battle_dir, f"banmen_{adjusted_tesuu}_resize.png")
            if os.path.exists(img_path):
                story.append(RLImage(img_path, width=7 * cm, height=11 * cm))
            story.append(Spacer(1, 0.3 * cm))
            continue

        if line.strip():
            story.append(Paragraph(xml_escape(line), body_style))
        else:
            story.append(Spacer(1, 0.3 * cm))

    # 棋譜データページ
    story.append(PageBreak())
    story.append(Paragraph("棋譜データ", title_style))
    story.append(Spacer(1, 0.3 * cm))
    for kif_line in all_kif_data.splitlines():
        text = kif_line if kif_line.strip() else " "
        story.append(Paragraph(xml_escape(text), kif_style))

    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    doc.build(story)
    print(f"PDF保存: {pdf_path}")


def save_local_report_json(battle_dir, battle_id, kif_data, response, ROOT, log_path, re_url):
    # ...
    report = {
        "各対局の振り返り": {
            battle_id: {
                "対局データ": kif_data,
                "将棋プロ棋士AIの解説": response,
                "解析URL": re_url  # ← これを追加
            }
        }
    }
    json_path = os.path.join(battle_dir, "local_report.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    print(f"JSON保存: {json_path}")


def _fetch_board_image(args):
    """get_final_board_image の並列実行ラッパー"""
    kif, tesuu, save_prefix, teban = args
    analysis_url = get_final_board_image(kif, save_prefix, teban)
    return tesuu, analysis_url


def run(ROOT, log_path):
    """メイン実行関数（対局は1局ずつ逐次、盤面図取得は最大5並列）"""
    print("=== Step2: Gemini局所分析 開始 ===")

    # PDF フォントをメインスレッドで事前登録（スレッドセーフでないため）
    _register_font(ROOT)

    base_dir = os.path.join(ROOT, "temp")

    for trg_folder in sorted(os.listdir(base_dir)):
        trg_path = os.path.join(base_dir, trg_folder)
        if not (os.path.isdir(trg_path) and trg_folder.startswith("trg_")):
            continue

        for entry in sorted(os.listdir(trg_path)):
            battle_dir = os.path.join(trg_path, entry)
            if not os.path.isdir(battle_dir):
                continue

            # 既に local_report.json がある場合はスキップ
            if os.path.exists(os.path.join(battle_dir, "local_report.json")):
                continue

            kif_path = os.path.join(battle_dir, "raw.kif")
            if not os.path.exists(kif_path):
                continue

            battle_id = os.path.basename(battle_dir)
            print(f"[*] 解析中: {battle_id}")

            try:
                with open(kif_path, "r", encoding="utf-8") as f:
                    raw_kif = f.read()

                # --- 経費削減ロジック: response.txt が既にあるかチェック ---
                response_path = os.path.join(battle_dir, "response.txt")
                if os.path.exists(response_path):
                    print(f"[*] 既存の response.txt を使用します: {battle_id}")
                    with open(response_path, "r", encoding="utf-8") as f:
                        response = f.read()
                else:
                    # Gemini 呼び出し（逐次：トークン制限対応）
                    response = call_gemini_analysis(raw_kif, log_path)
                    print(response)
                    with open(response_path, "w", encoding="utf-8") as f:
                        f.write(response)

                # 貼り付け手数と棋譜データを取得
                haru_tesuu, kif_list, teban = txt_to_docx_with_images(raw_kif, response)

                # 盤面図を最大5並列で取得
                tasks = [
                    (kif, tesuu, os.path.join(battle_dir, f"banmen_{tesuu}"), teban)
                    for kif, tesuu in zip(kif_list, haru_tesuu)
                ]
                re_url = ""
                with ThreadPoolExecutor(max_workers=5) as executor:
                    for tesuu, analysis_url in executor.map(_fetch_board_image, tasks):
                        if tesuu == "all":
                            re_url = analysis_url

                # docx 作成
                docx_path = os.path.join(battle_dir, "local_report.docx")
                make_leport_docx(response, battle_id, re_url, battle_dir, docx_path, raw_kif, haru_tesuu)

                # PDF 生成
                pdf_path = os.path.join(battle_dir, "local_report.pdf")
                make_report_pdf(ROOT,response, battle_id, re_url, battle_dir, pdf_path, raw_kif, haru_tesuu)

                entity_kif = cut_unnecessary_kif(ROOT, log_path, kif_path)

                # local_report.json 保存
                save_local_report_json(battle_dir, battle_id, entity_kif, response, ROOT, log_path, re_url)

                print(f"[OK] 完了: {battle_id}")

                # --- クリーンアップ: 不要なファイルを削除 ---
                delete_png_and_docx_files(battle_dir)

            except Exception as e:
                error_msg = f"[ERROR] {battle_id}: {str(e)}\n{traceback.format_exc()}\n"
                print(error_msg)
                _write_log(log_path, error_msg)

    print("=== Step2: Gemini局所分析 完了 ===")